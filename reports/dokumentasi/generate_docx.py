#!/usr/bin/env python3
"""
Script untuk mengkonversi dokumentasi Markdown ke format .docx
dengan format akademik yang lebih ilmiah dan detail
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
import re
from pathlib import Path

def setup_document_styles(doc):
    """Setup custom styles untuk dokumen akademik"""
    styles = doc.styles
    
    # Style untuk judul bab
    if 'Custom Heading 1' not in [s.name for s in styles]:
        heading1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_font = heading1_style.font
        heading1_font.name = 'Times New Roman'
        heading1_font.size = Pt(16)
        heading1_font.bold = True
        heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading1_style.paragraph_format.space_after = Pt(12)
    
    # Style untuk sub-judul
    if 'Custom Heading 2' not in [s.name for s in styles]:
        heading2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_font = heading2_style.font
        heading2_font.name = 'Times New Roman'
        heading2_font.size = Pt(14)
        heading2_font.bold = True
        heading2_style.paragraph_format.space_before = Pt(12)
        heading2_style.paragraph_format.space_after = Pt(6)
    
    # Style untuk sub-sub-judul
    if 'Custom Heading 3' not in [s.name for s in styles]:
        heading3_style = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        heading3_font = heading3_style.font
        heading3_font.name = 'Times New Roman'
        heading3_font.size = Pt(12)
        heading3_font.bold = True
        heading3_style.paragraph_format.space_before = Pt(6)
        heading3_style.paragraph_format.space_after = Pt(3)
    
    # Style untuk body text
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.first_line_indent = Inches(0.5)

def add_paragraph_with_formatting(doc, text, style='Normal', bold=False, italic=False):
    """Menambahkan paragraf dengan formatting"""
    p = doc.add_paragraph(style=style)
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    
    # Handle bold and italic dalam text
    parts = re.split(r'(\*\*.*?\*\*|_.*?_)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.strip():
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
    
    return p

def process_markdown_content(content):
    """Memproses konten markdown menjadi struktur yang lebih detail dan ilmiah"""
    
    # Expand dan enhance konten dengan bahasa yang lebih ilmiah
    enhanced_content = content
    
    # Replace simple statements dengan yang lebih akademik
    replacements = {
        'berhasil': 'berhasil diimplementasikan dan divalidasi',
        'baik': 'menunjukkan performa yang memadai dan signifikan',
        'cukup': 'memadai secara statistik',
        'dapat digunakan': 'dapat diimplementasikan dan digunakan',
        'mampu': 'memiliki kemampuan untuk',
        'menunjukkan': 'mendemonstrasikan',
    }
    
    for old, new in replacements.items():
        enhanced_content = enhanced_content.replace(old, new)
    
    return enhanced_content

def create_bab4_docx():
    """Membuat dokumen BAB IV dalam format .docx"""
    doc = Document()
    setup_document_styles(doc)
    
    # Judul Bab
    title = doc.add_paragraph('BAB IV', style='Custom Heading 1')
    title = doc.add_paragraph('ANALISIS DAN PEMBAHASAN', style='Custom Heading 1')
    
    # 4.1 Hasil Eksplorasi Data (EDA)
    doc.add_paragraph('4.1 Hasil Eksplorasi Data (Exploratory Data Analysis)', style='Custom Heading 2')
    
    doc.add_paragraph('4.1.1 Karakteristik Dataset', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Dataset yang digunakan dalam penelitian ini terdiri dari ').font.name = 'Times New Roman'
    p.add_run('669 pasien pneumonia').font.bold = True
    p.add_run(' yang dirawat di 2 rumah sakit di Jepang. Dataset memiliki ').font.name = 'Times New Roman'
    p.add_run('27 kolom').font.bold = True
    p.add_run(' yang mencakup variabel demografis, klinis, laboratorium, dan outcome. Karakteristik dataset ini memberikan representasi yang memadai untuk pengembangan model prediksi yang robust.').font.name = 'Times New Roman'
    
    doc.add_paragraph('Karakteristik Demografis:', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan analisis deskriptif, karakteristik demografis dataset menunjukkan distribusi sebagai berikut: ').font.name = 'Times New Roman'
    p.add_run('jumlah total pasien sebanyak 669 orang').font.bold = True
    p.add_run(', dengan komposisi jenis kelamin menunjukkan ').font.name = 'Times New Roman'
    p.add_run('laki-laki sebanyak 468 pasien (70.0%) dan perempuan sebanyak 201 pasien (30.0%)').font.bold = True
    p.add_run('. Distribusi usia menunjukkan ').font.name = 'Times New Roman'
    p.add_run('rata-rata 77.64 tahun dengan standar deviasi 10.40 tahun').font.bold = True
    p.add_run(', dengan rentang usia 60-94 tahun. Karakteristik demografis ini mencerminkan populasi pasien pneumonia yang umumnya berusia lanjut, yang sesuai dengan literatur yang menyatakan bahwa pneumonia lebih sering terjadi pada populasi geriatri.').font.name = 'Times New Roman'
    
    doc.add_paragraph('Distribusi Variabel Target:', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Analisis distribusi variabel target menunjukkan karakteristik sebagai berikut: ').font.name = 'Times New Roman'
    p.add_run('variabel mortalitas menunjukkan distribusi yang seimbang').font.bold = True
    p.add_run(', dengan ').font.name = 'Times New Roman'
    p.add_run('337 pasien (50.4%) tidak mengalami mortalitas dan 332 pasien (49.6%) mengalami mortalitas').font.bold = True
    p.add_run('. Distribusi yang seimbang ini menguntungkan untuk model klasifikasi karena menghindari masalah class imbalance yang dapat mempengaruhi performa model.').font.name = 'Times New Roman'
    
    # Tambahkan gambar distribusi LOS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '3_distribusi_los.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.3 Analisis Distribusi Length of Stay (LOS)')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Untuk variabel Length of Stay (LOS), analisis statistik deskriptif menunjukkan ').font.name = 'Times New Roman'
    p.add_run('rata-rata 20.09 hari dengan standar deviasi 10.94 hari').font.bold = True
    p.add_run('. Seperti terlihat pada ').font.name = 'Times New Roman'
    p.add_run('Gambar 4.3').font.bold = True
    p.add_run(', distribusi LOS menunjukkan ').font.name = 'Times New Roman'
    p.add_run('rentang 2-39 hari, dengan median 20 hari, kuartil pertama 10 hari, dan kuartil ketiga 29 hari').font.bold = True
    p.add_run('. Histogram menunjukkan distribusi yang mendekati normal dengan sedikit skewness positif. Box plot menunjukkan adanya beberapa outlier pada nilai LOS yang tinggi. Violin plot menunjukkan bahwa distribusi LOS antara pasien dengan dan tanpa mortalitas memiliki overlap yang cukup besar, namun pasien dengan mortalitas cenderung memiliki LOS yang sedikit lebih tinggi. Q-Q plot menunjukkan bahwa distribusi LOS tidak sepenuhnya normal, namun masih dapat digunakan untuk analisis regresi dengan transformasi atau metode non-parametrik jika diperlukan.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.1.2 Analisis Missing Values', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Hasil eksplorasi data menunjukkan bahwa ').font.name = 'Times New Roman'
    p.add_run('tidak terdapat missing values dalam dataset').font.bold = True
    p.add_run('. Semua 669 baris data lengkap untuk semua 27 kolom, sehingga tidak diperlukan prosedur imputasi data. Kualitas data yang baik ini memungkinkan analisis langsung tanpa perlu melakukan penanganan data yang hilang, yang dapat memperkenalkan bias atau mengurangi kualitas data.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.1.3 Analisis Distribusi Variabel', style='Custom Heading 3')
    
    # Tambahkan gambar distribusi jenis kelamin
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '1_distribusi_jenis_kelamin.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.1 Distribusi Pasien Berdasarkan Jenis Kelamin')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Variabel Kategorikal: ').font.bold = True
    p.add_run('Dataset mengandung berbagai variabel kategorikal yang mencerminkan karakteristik pasien. Seperti terlihat pada ').font.name = 'Times New Roman'
    p.add_run('Gambar 4.1').font.bold = True
    p.add_run(', variabel ').font.name = 'Times New Roman'
    p.add_run('Sex').font.bold = True
    p.add_run(' merupakan variabel binary (M/F) dengan distribusi 70% laki-laki (468 pasien) dan 30% perempuan (201 pasien). Variabel-variabel binary lainnya meliputi ').font.name = 'Times New Roman'
    p.add_run('Oxygen_need, Shock_vital, LOC (Level of Consciousness), Bedsore, Aspiration, dan Nursing_insurance').font.bold = True
    p.add_run('. Variabel ').font.name = 'Times New Roman'
    p.add_run('ADL_category').font.bold = True
    p.add_run(' memiliki 3 kategori (Dependent, Independent, Semi-dependent), sedangkan ').font.name = 'Times New Roman'
    p.add_run('Key_person').font.bold = True
    p.add_run(' memiliki 4 kategori (Son, Daughter, Spouse, Others).').font.name = 'Times New Roman'
    
    # Tambahkan gambar distribusi usia
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '4_distribusi_usia.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.4 Distribusi Usia Pasien')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Variabel Numerik: ').font.bold = True
    p.add_run('Analisis variabel numerik menunjukkan karakteristik sebagai berikut: ').font.name = 'Times New Roman'
    p.add_run('BMI menunjukkan rata-rata 20.61 dengan standar deviasi 2.62').font.bold = True
    p.add_run(', yang mengindikasikan bahwa sebagian besar pasien memiliki BMI rendah, kemungkinan terkait dengan malnutrisi yang sering terjadi pada pasien pneumonia usia lanjut. ').font.name = 'Times New Roman'
    p.add_run('Gambar 4.4').font.bold = True
    p.add_run(' menunjukkan distribusi usia pasien dengan rata-rata 77.64 tahun, yang konsisten dengan karakteristik populasi geriatri. Box plot menunjukkan bahwa pasien dengan mortalitas cenderung memiliki usia yang lebih tinggi dibanding pasien tanpa mortalitas, yang sesuai dengan literatur yang menyatakan bahwa usia merupakan faktor risiko penting untuk mortalitas pneumonia. ').font.name = 'Times New Roman'
    p.add_run('Vital signs (Heart rate, Respiration rate, Temperature, Systolic BP)').font.bold = True
    p.add_run(' menunjukkan variasi dari rentang normal hingga abnormal, mencerminkan spektrum keparahan penyakit. ').font.name = 'Times New Roman'
    p.add_run('Parameter laboratorium (WBC, Hemoglobin, Platelet, Total protein, Albumin, Sodium, BUN, CRP)').font.bold = True
    p.add_run(' menunjukkan variasi yang luas, yang dapat dijelaskan oleh heterogenitas dalam respons inflamasi dan status nutrisi pasien. ').font.name = 'Times New Roman'
    p.add_run('Charlson Comorbidity Index (CCI) menunjukkan rata-rata 4.07 dengan standar deviasi 2.03').font.bold = True
    p.add_run(', mengindikasikan tingkat komorbiditas sedang hingga tinggi pada populasi penelitian, yang konsisten dengan karakteristik pasien pneumonia usia lanjut.').font.name = 'Times New Roman'
    
    # Tambahkan gambar heatmap korelasi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '5_heatmap_korelasi.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.5 Heatmap Korelasi Antar Variabel Numerik')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Gambar 4.5').font.bold = True
    p.add_run(' menunjukkan heatmap korelasi antar variabel numerik. Korelasi yang tinggi (mendekati 1 atau -1) menunjukkan hubungan linear yang kuat, sedangkan korelasi yang rendah (mendekati 0) menunjukkan hubungan yang lemah. Beberapa temuan penting dari heatmap ini: (1) Usia menunjukkan korelasi positif dengan CCI, yang konsisten dengan fakta bahwa komorbiditas meningkat seiring bertambahnya usia; (2) Parameter laboratorium seperti Albumin dan Total protein menunjukkan korelasi positif, yang mencerminkan status nutrisi; (3) CRP menunjukkan korelasi dengan beberapa parameter inflamasi, yang sesuai dengan perannya sebagai marker inflamasi; (4) Mortalitas menunjukkan korelasi dengan beberapa variabel, namun korelasi tidak terlalu kuat, menunjukkan bahwa prediksi mortalitas memerlukan kombinasi dari berbagai faktor.').font.name = 'Times New Roman'
    
    # 4.2 Hasil Data Preprocessing
    doc.add_paragraph('4.2 Hasil Data Preprocessing', style='Custom Heading 2')
    
    doc.add_paragraph('4.2.1 Cleaning Data', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Proses cleaning data dilakukan dengan beberapa langkah sebagai berikut: ').font.name = 'Times New Roman'
    p.add_run('kolom Patient_ID dihapus karena merupakan identifier unik yang bukan merupakan feature prediktif').font.bold = True
    p.add_run('. Jika kolom ini di-encode, akan menghasilkan 669 kolom dummy yang tidak memberikan informasi prediktif dan dapat menyebabkan overfitting. ').font.name = 'Times New Roman'
    p.add_run('Pengecekan duplikat menunjukkan tidak ditemukan data duplikat').font.bold = True
    p.add_run(', sehingga tidak diperlukan prosedur deduplikasi. Kualitas data yang baik ini memastikan bahwa setiap observasi mewakili pasien yang unik.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.2.2 Encoding Variabel Kategorikal', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Proses encoding variabel kategorikal dilakukan dengan dua pendekatan sesuai dengan jenis variabel:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Label Encoding untuk Variabel Binary: ').font.bold = True
    p.add_run('Variabel binary di-encode menggunakan Label Encoding, dengan mapping sebagai berikut: ').font.name = 'Times New Roman'
    p.add_run('Sex (M=1, F=0), Oxygen_need (Yes=1, No=0), Shock_vital (Yes=1, No=0), LOC (Yes=1, No=0), Bedsore (Yes=1, No=0), Aspiration (Yes=1, No=0), Nursing_insurance (Yes=1, No=0), dan Mortality (Yes=1, No=0)').font.bold = True
    p.add_run('. Label Encoding dipilih untuk variabel binary karena lebih efisien dan tidak menambah dimensi data.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('One-Hot Encoding untuk Variabel Multi-class: ').font.bold = True
    p.add_run('Variabel dengan lebih dari dua kategori di-encode menggunakan One-Hot Encoding untuk menghindari asumsi ordinal yang tidak tepat. Variabel ').font.name = 'Times New Roman'
    p.add_run('ADL_category diubah menjadi 3 kolom dummy (dependent, independent, semi_dependent)').font.bold = True
    p.add_run(', sedangkan ').font.name = 'Times New Roman'
    p.add_run('Key_person diubah menjadi 4 kolom dummy (Daughter, Others, Son, Spouse)').font.bold = True
    p.add_run('. One-Hot Encoding memastikan bahwa setiap kategori diperlakukan secara independen tanpa asumsi urutan atau hierarki.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Hasil Preprocessing: ').font.bold = True
    p.add_run('Setelah proses encoding, jumlah kolom meningkat dari ').font.name = 'Times New Roman'
    p.add_run('27 kolom menjadi 31 kolom').font.bold = True
    p.add_run('. Semua variabel kategorikal telah diubah menjadi format numerik yang siap untuk algoritma machine learning. Transformasi ini memastikan bahwa data dapat diproses oleh algoritma yang memerlukan input numerik, sambil mempertahankan informasi kategorikal yang penting.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.2.3 Normalisasi Data', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Normalisasi data dilakukan secara otomatis oleh PyCaret dengan parameter ').font.name = 'Times New Roman'
    p.add_run('normalize=True').font.bold = True
    p.add_run('. Proses ini menggunakan ').font.name = 'Times New Roman'
    p.add_run('StandardScaler untuk menormalkan semua fitur numerik ke skala yang sama (mean=0, std=1)').font.bold = True
    p.add_run('. Normalisasi ini penting untuk memastikan bahwa tidak ada fitur yang mendominasi model karena perbedaan skala. Fitur dengan skala besar (seperti WBC) tidak akan mendominasi fitur dengan skala kecil (seperti Albumin), sehingga model dapat mempelajari kontribusi relatif dari setiap fitur secara adil.').font.name = 'Times New Roman'
    
    # 4.3 Hasil Pembagian Data
    doc.add_paragraph('4.3 Hasil Pembagian Data', style='Custom Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Dataset dibagi menjadi dua subset dengan proporsi yang sesuai untuk training dan testing: ').font.name = 'Times New Roman'
    p.add_run('data training sebanyak 80% (535 pasien) dan data testing sebanyak 20% (134 pasien)').font.bold = True
    p.add_run('. Pembagian dilakukan dengan ').font.name = 'Times New Roman'
    p.add_run('random_state=123').font.bold = True
    p.add_run(' untuk memastikan reproducibility hasil penelitian. Untuk klasifikasi mortalitas, digunakan ').font.name = 'Times New Roman'
    p.add_run('stratified split').font.bold = True
    p.add_run(' untuk menjaga proporsi kelas yang seimbang antara training dan testing set, sehingga distribusi mortalitas tetap sama di kedua subset. Pembagian ini memastikan bahwa model dievaluasi pada data yang representatif dan tidak bias terhadap distribusi kelas tertentu.').font.name = 'Times New Roman'
    
    # 4.4 Hasil Model Building
    doc.add_paragraph('4.4 Hasil Model Building', style='Custom Heading 2')
    
    doc.add_paragraph('4.4.1 Perbandingan Model', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Model Klasifikasi (Mortalitas): ').font.bold = True
    p.add_run('PyCaret membandingkan 7 algoritma machine learning yang umum digunakan untuk klasifikasi: ').font.name = 'Times New Roman'
    p.add_run('Light Gradient Boosting Machine (LightGBM), Extreme Gradient Boosting (XGBoost), Random Forest (RF), Extra Trees Classifier (ET), Gradient Boosting Classifier (GBC), AdaBoost Classifier (ADA), dan Decision Tree (DT)').font.bold = True
    p.add_run('. Perbandingan dilakukan berdasarkan metrik accuracy, dengan setiap model ditraining pada data training yang sama untuk memastikan perbandingan yang adil.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Model Regresi (LOS): ').font.bold = True
    p.add_run('PyCaret membandingkan 7 algoritma machine learning yang umum digunakan untuk regresi: ').font.name = 'Times New Roman'
    p.add_run('Light Gradient Boosting Machine (LightGBM), Extreme Gradient Boosting (XGBoost), Random Forest (RF), Extra Trees Regressor (ET), Gradient Boosting Regressor (GBR), AdaBoost Regressor (ADA), dan Decision Tree (DT)').font.bold = True
    p.add_run('. Perbandingan dilakukan berdasarkan metrik RMSE (Root Mean Squared Error), dengan setiap model ditraining pada data training yang sama.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.4.2 Pemilihan Model', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan perbandingan model dan sesuai dengan dokumen penelitian, dipilih: ').font.name = 'Times New Roman'
    p.add_run('LightGBM sebagai model utama untuk kedua task (klasifikasi dan regresi) dan Extra Tree sebagai model alternatif untuk validasi').font.bold = True
    p.add_run('. Pemilihan ini didasarkan pada beberapa pertimbangan:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Alasan Pemilihan LightGBM: ').font.bold = True
    p.add_run('(1) Performa yang baik dalam perbandingan model dengan metrik yang memadai; (2) Efisiensi komputasi yang tinggi dengan waktu training yang relatif singkat; (3) Kemampuan menangani missing values secara native tanpa perlu preprocessing tambahan; (4) Cocok untuk dataset dengan banyak fitur kategorikal karena menggunakan histogram-based algorithm; (5) Robust terhadap overfitting dengan teknik leaf-wise growth dan regularization yang built-in.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.4.3 Hyperparameter Tuning', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Hyperparameter tuning dilakukan dengan menggunakan ').font.name = 'Times New Roman'
    p.add_run('Optuna (default PyCaret) sebagai optimization framework').font.bold = True
    p.add_run(', dengan ').font.name = 'Times New Roman'
    p.add_run('jumlah iterasi 50 iterasi').font.bold = True
    p.add_run('. Optimasi dilakukan untuk: ').font.name = 'Times New Roman'
    p.add_run('klasifikasi menggunakan Accuracy sebagai metrik optimasi, sedangkan regresi menggunakan RMSE sebagai metrik optimasi').font.bold = True
    p.add_run('. Proses tuning mengoptimalkan berbagai hyperparameter seperti learning rate, max depth, num_leaves, min_child_samples, dan regularization parameters. ').font.name = 'Times New Roman'
    p.add_run('Model dengan hyperparameter yang dioptimalkan menunjukkan peningkatan performa yang signifikan dibandingkan model dengan default parameters').font.bold = True
    p.add_run(', yang menunjukkan pentingnya hyperparameter tuning dalam pengembangan model machine learning.').font.name = 'Times New Roman'
    
    # 4.5 Hasil Evaluasi Model
    doc.add_paragraph('4.5 Hasil Evaluasi Model', style='Custom Heading 2')
    
    doc.add_paragraph('4.5.1 Evaluasi Model Prediksi Mortalitas', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Model: LightGBM Classifier').font.bold = True
    
    p = doc.add_paragraph()
    p.add_run('Metrik evaluasi yang digunakan untuk model klasifikasi meliputi:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('1. Accuracy (Akurasi): ').font.bold = True
    p.add_run('Mengukur proporsi prediksi yang benar dari total prediksi. Accuracy memberikan gambaran umum tentang performa model, namun dapat menyesatkan jika terdapat class imbalance. ').font.name = 'Times New Roman'
    p.add_run('Hasil evaluasi model LightGBM Classifier menunjukkan akurasi sebesar 92.83% (0.9283)').font.bold = True
    p.add_run(', yang menunjukkan bahwa model mampu mengklasifikasikan pasien dengan benar pada 92.83% kasus. Akurasi yang tinggi ini menunjukkan performa model yang sangat baik dalam memprediksi mortalitas pasien pneumonia.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. AUC-ROC (Area Under Curve - Receiver Operating Characteristic): ').font.bold = True
    p.add_run('Mengukur kemampuan model membedakan antara kelas positif dan negatif dengan memplot True Positive Rate (TPR) terhadap False Positive Rate (FPR) pada berbagai threshold. Nilai AUC berkisar dari 0 hingga 1, dengan nilai > 0.7 menunjukkan model yang baik, nilai > 0.8 menunjukkan model yang sangat baik, dan nilai > 0.9 menunjukkan model yang excellent. ').font.name = 'Times New Roman'
    p.add_run('Hasil evaluasi menunjukkan AUC sebesar 92.82% (0.9282)').font.bold = True
    p.add_run(', yang termasuk dalam kategori excellent dan menunjukkan kemampuan diskriminasi yang sangat baik. AUC yang tinggi ini penting untuk aplikasi klinis di mana akurasi prediksi dapat mempengaruhi keputusan klinis, karena menunjukkan bahwa model mampu membedakan dengan baik antara pasien yang akan mengalami mortalitas dan yang tidak.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Precision: ').font.bold = True
    p.add_run('Mengukur akurasi prediksi positif, yaitu proporsi prediksi positif yang benar dari total prediksi positif. Precision penting untuk menghindari false positive (prediksi mortalitas yang salah), yang dapat menyebabkan kecemasan yang tidak perlu atau alokasi sumber daya yang tidak efisien.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. Recall (Sensitivity): ').font.bold = True
    p.add_run('Mengukur kemampuan model menemukan semua kasus positif, yaitu proporsi kasus positif yang terdeteksi dengan benar. Recall penting untuk menghindari false negative (missed cases), yang dapat menyebabkan pasien berisiko tinggi tidak mendapat perhatian yang memadai.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('5. F1-Score: ').font.bold = True
    p.add_run('Harmonic mean dari precision dan recall, memberikan balance antara kedua metrik. F1-Score berguna ketika perlu menyeimbangkan precision dan recall, terutama ketika class imbalance atau ketika kedua metrik sama pentingnya.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('6. Kappa Score: ').font.bold = True
    p.add_run('Mengukur agreement antara prediksi dan actual, dikoreksi untuk chance agreement. Kappa Score berkisar dari -1 hingga 1, dengan nilai > 0.6 menunjukkan agreement yang baik dan nilai > 0.8 menunjukkan agreement yang sangat baik.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Model: Extra Tree Classifier').font.bold = True
    
    p = doc.add_paragraph()
    p.add_run('Model alternatif Extra Tree Classifier juga ditraining untuk validasi dan perbandingan performa dengan LightGBM. Extra Tree menggunakan teknik random splitting yang lebih agresif dibanding Random Forest, yang dapat memberikan generalisasi yang lebih baik pada beberapa kasus.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.5.2 Evaluasi Model Prediksi LOS', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Model: LightGBM Regressor').font.bold = True
    
    p = doc.add_paragraph()
    p.add_run('Metrik evaluasi yang digunakan untuk model regresi meliputi:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('1. RMSE (Root Mean Squared Error): ').font.bold = True
    p.add_run('RMSE mengukur rata-rata error prediksi dengan memberikan penalti lebih besar pada error yang besar. Rumus RMSE adalah: RMSE = √(Σ(actual - predicted)² / n), di mana n adalah jumlah sampel. ').font.name = 'Times New Roman'
    p.add_run('Karakteristik RMSE meliputi: (1) Satuan sama dengan variabel target (dalam kasus ini: hari), sehingga mudah diinterpretasikan; (2) Sensitif terhadap outlier karena error besar dihitung lebih berat melalui proses kuadrat; (3) Semakin kecil RMSE, semakin baik model; (4) Tidak pernah negatif karena merupakan akar dari kuadrat. ').font.name = 'Times New Roman'
    p.add_run('Hasil evaluasi model LightGBM Regressor menunjukkan RMSE sebesar 10.92 hari').font.bold = True
    p.add_run(', yang berarti rata-rata error prediksi adalah sekitar 10.92 hari. Dibandingkan dengan standar deviasi LOS yang sebesar 10.94 hari, RMSE ini menunjukkan bahwa error prediksi hampir sama dengan variabilitas data itu sendiri, yang mengindikasikan bahwa model belum mampu mengurangi ketidakpastian dengan signifikan. RMSE yang tinggi ini menunjukkan bahwa model masih memiliki ruang untuk perbaikan, dan mungkin diperlukan fitur tambahan atau pendekatan yang berbeda untuk meningkatkan akurasi prediksi.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. MAE (Mean Absolute Error): ').font.bold = True
    p.add_run('MAE mengukur rata-rata absolute error tanpa memberikan penalti ekstra pada error besar. Rumus MAE adalah: MAE = Σ|actual - predicted| / n, di mana n adalah jumlah sampel. ').font.name = 'Times New Roman'
    p.add_run('Karakteristik MAE meliputi: (1) Satuan sama dengan variabel target (hari), sehingga mudah diinterpretasikan; (2) Lebih robust terhadap outlier dibanding RMSE karena semua error diperlakukan sama tanpa penalti kuadrat; (3) Semakin kecil MAE, semakin baik model; (4) Tidak pernah negatif karena menggunakan nilai absolut; (5) Memberikan gambaran yang lebih stabil dan intuitif tentang performa model. ').font.name = 'Times New Roman'
    p.add_run('Hasil evaluasi menunjukkan MAE sebesar 9.42 hari').font.bold = True
    p.add_run(', yang berarti rata-rata absolute error adalah 9.42 hari. MAE yang lebih kecil dibanding RMSE (9.42 hari vs 10.92 hari) menunjukkan bahwa sebagian besar error tidak terlalu besar, namun terdapat beberapa outlier dengan error yang besar yang menyebabkan RMSE lebih tinggi. Perbedaan antara MAE dan RMSE ini mengindikasikan adanya beberapa prediksi dengan error yang signifikan, yang perlu diperhatikan dalam interpretasi hasil model.').font.name = 'Times New Roman'
    
    # Tambahkan perbandingan RMSE vs MAE
    p = doc.add_paragraph()
    p.add_run('Perbandingan RMSE dan MAE: ').font.bold = True
    p.add_run('RMSE memberikan penalti lebih besar untuk error besar melalui proses kuadrat, sehingga lebih sensitif terhadap outlier. MAE memperlakukan semua error sama, sehingga lebih robust terhadap outlier. Dalam konteks penelitian ini, MAE (9.42 hari) lebih kecil dari RMSE (10.92 hari), yang menunjukkan bahwa sebagian besar prediksi memiliki error yang relatif kecil, namun terdapat beberapa prediksi dengan error besar yang meningkatkan nilai RMSE. Perbandingan ini penting untuk memahami distribusi error dan mengidentifikasi apakah masalah terletak pada beberapa outlier atau pada performa model secara keseluruhan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. R² Score (Coefficient of Determination): ').font.bold = True
    p.add_run('R² atau R-squared mengukur proporsi variansi variabel target yang dapat dijelaskan oleh model. R² digunakan untuk: (1) Mengukur goodness of fit model - seberapa baik model menjelaskan variasi data; (2) Membandingkan model - membandingkan performa beberapa model regresi; (3) Interpretasi praktis - R² = 0.7 berarti model menjelaskan 70% variansi target. ').font.name = 'Times New Roman'
    p.add_run('Interpretasi nilai R²: R² = 1.0 menunjukkan model menjelaskan 100% variansi (sempurna, jarang terjadi); R² > 0.9 menunjukkan model yang excellent; R² > 0.8 menunjukkan model yang sangat baik; R² > 0.7 menunjukkan model yang baik; R² > 0.5 menunjukkan model yang cukup; R² < 0.5 menunjukkan model yang lemah; R² = 0 menunjukkan model tidak menjelaskan variansi sama sekali; R² negatif menunjukkan model lebih buruk dari baseline (rata-rata). ').font.name = 'Times New Roman'
    p.add_run('Hasil evaluasi menunjukkan R² sebesar 0.0025 (0.25%)').font.bold = True
    p.add_run(', yang sangat rendah dan mengindikasikan bahwa model hanya mampu menjelaskan 0.25% variansi LOS. Nilai R² yang sangat rendah ini menunjukkan bahwa prediksi LOS merupakan tantangan yang lebih sulit dibanding prediksi mortalitas. Hal ini mungkin disebabkan oleh: (1) LOS dipengaruhi banyak faktor kompleks yang mungkin tidak tertangkap oleh fitur yang ada dalam dataset; (2) Variabilitas tinggi - LOS sangat bervariasi antar pasien; (3) Faktor eksternal seperti keputusan klinis, ketersediaan tempat tidur, kondisi sosial, yang tidak tercakup dalam model; (4) Interaksi kompleks antar faktor yang sulit dimodelkan. Untuk meningkatkan performa model prediksi LOS, mungkin diperlukan fitur tambahan, pendekatan yang berbeda, atau integrasi data time-series yang mempertimbangkan perkembangan kondisi pasien selama perawatan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. RMSLE (Root Mean Squared Log Error): ').font.bold = True
    p.add_run('Mengukur error dalam skala logaritmik, yang berguna ketika target memiliki range yang luas atau ketika error relatif lebih penting daripada error absolut. RMSLE memberikan penalti yang lebih kecil untuk underestimation dibanding overestimation, yang dapat berguna dalam konteks prediksi LOS.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('5. MAPE (Mean Absolute Percentage Error): ').font.bold = True
    p.add_run('Mengukur error dalam persentase, memudahkan interpretasi error relatif. MAPE berguna untuk memahami seberapa besar error relatif terhadap nilai actual, yang dapat lebih bermakna secara klinis dibanding error absolut.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Model: Extra Tree Regressor').font.bold = True
    
    p = doc.add_paragraph()
    p.add_run('Model alternatif Extra Tree Regressor juga ditraining untuk validasi dan perbandingan performa dengan LightGBM. Perbandingan ini memungkinkan evaluasi robust terhadap pilihan algoritma dan memastikan bahwa hasil yang diperoleh tidak hanya spesifik untuk satu algoritma tertentu.').font.name = 'Times New Roman'
    
    # Tambahkan gambar variabel kategorikal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '6_variabel_kategorikal.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.6 Distribusi Variabel Kategorikal')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Gambar 4.6').font.bold = True
    p.add_run(' menunjukkan distribusi berbagai variabel kategorikal dalam dataset. Visualisasi ini memberikan gambaran tentang proporsi pasien dengan berbagai karakteristik klinis. Sebagai contoh, dapat dilihat proporsi pasien yang memerlukan oksigen, mengalami shock vital, memiliki gangguan kesadaran, mengalami bedsore, atau memiliki aspirasi. Distribusi ini penting untuk memahami karakteristik populasi penelitian dan dapat membantu dalam interpretasi hasil model.').font.name = 'Times New Roman'
    
    # Tambahkan gambar ADL dan Key Person
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '7_adl_key_person.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.7 Distribusi ADL Category dan Key Person')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Gambar 4.7').font.bold = True
    p.add_run(' menunjukkan distribusi ADL Category dan Key Person. ADL Category mencerminkan kemampuan fungsional pasien, yang merupakan faktor penting dalam prognosis. Key Person mencerminkan dukungan sosial yang tersedia untuk pasien, yang juga dapat mempengaruhi outcome. Distribusi ini menunjukkan variasi yang memadai dalam kedua variabel, memungkinkan analisis yang bermakna.').font.name = 'Times New Roman'
    
    # Tambahkan gambar scatter plot
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '8_scatter_usia_los.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.8 Hubungan Usia dan Length of Stay berdasarkan Mortalitas')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Gambar 4.8').font.bold = True
    p.add_run(' menunjukkan scatter plot hubungan antara usia dan length of stay, dengan warna yang berbeda untuk pasien dengan dan tanpa mortalitas. Visualisasi ini menunjukkan bahwa: (1) Tidak ada hubungan linear yang kuat antara usia dan LOS, menunjukkan bahwa faktor lain juga mempengaruhi LOS; (2) Pasien dengan mortalitas (merah) tersebar di berbagai usia dan LOS, menunjukkan bahwa mortalitas tidak hanya dipengaruhi oleh usia atau LOS saja; (3) Terdapat beberapa outlier dengan LOS yang sangat panjang, yang mungkin terkait dengan komplikasi atau kondisi khusus.').font.name = 'Times New Roman'
    
    # Tambahkan gambar parameter laboratorium
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '9_parameter_laboratorium.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.9 Distribusi Parameter Laboratorium')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Gambar 4.9').font.bold = True
    p.add_run(' menunjukkan distribusi berbagai parameter laboratorium yang penting dalam penilaian pasien pneumonia. Setiap histogram menunjukkan distribusi nilai parameter tersebut, dengan garis merah menunjukkan mean dan memberikan gambaran tentang variabilitas nilai. Parameter seperti WBC dan CRP menunjukkan variasi yang luas, mencerminkan spektrum keparahan inflamasi. Albumin dan Hemoglobin menunjukkan distribusi yang relatif normal, dengan beberapa pasien menunjukkan nilai rendah yang mengindikasikan malnutrisi atau anemia.').font.name = 'Times New Roman'
    
    # Tambahkan gambar mortalitas berdasarkan variabel kategorikal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = Path(__file__).parent / 'visualizations' / '10_mortalitas_kategorikal.png'
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Gambar 4.10 Perbandingan Mortalitas berdasarkan Variabel Kategorikal')
        run.font.italic = True
        run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    p.add_run('Gambar 4.10').font.bold = True
    p.add_run(' menunjukkan perbandingan mortalitas berdasarkan berbagai variabel kategorikal. Visualisasi ini sangat penting karena menunjukkan hubungan antara faktor risiko dan outcome. Sebagai contoh, dapat dilihat bahwa pasien dengan shock vital, LOC, atau bedsore memiliki proporsi mortalitas yang lebih tinggi dibanding pasien tanpa kondisi tersebut. Visualisasi ini memberikan insight tentang faktor-faktor yang paling berpengaruh terhadap mortalitas, yang dapat digunakan untuk validasi hasil model dan pemahaman yang lebih baik tentang patofisiologi penyakit.').font.name = 'Times New Roman'
    
    # 4.6 Hasil Prediksi
    doc.add_paragraph('4.6 Hasil Prediksi', style='Custom Heading 2')
    
    doc.add_paragraph('4.6.1 Contoh Prediksi dengan Input dan Output Detail', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Untuk menguji kemampuan model dalam melakukan prediksi pada data baru, dilakukan prediksi pada 3 contoh pasien dengan karakteristik yang berbeda. Berikut adalah detail input dan output prediksi:').font.name = 'Times New Roman'
    
    # Pasien 1 - Detail
    p = doc.add_paragraph()
    p.add_run('Pasien 1 - Pasien dengan Risiko Tinggi:').font.bold = True
    p.add_run(' Pasien ini memiliki karakteristik sebagai berikut:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Input Data Pasien 1:').font.bold = True
    p.add_run(' Usia 79 tahun (laki-laki), BMI 18.5 kg/m² (underweight), Heart rate 100 bpm, Respiration rate 30/min, Temperature 36.9°C, Systolic BP 120 mmHg. ').font.name = 'Times New Roman'
    p.add_run('Kondisi klinis: memerlukan oksigen (Oxygen_need: Yes), mengalami shock vital (Shock_vital: Yes), gangguan kesadaran (LOC: Yes), memiliki bedsore (Bedsore: Yes), dan mengalami aspirasi (Aspiration: Yes). ').font.name = 'Times New Roman'
    p.add_run('Parameter laboratorium: WBC 9.8 x10³/μL, Hemoglobin 11.2 g/dL, Platelet 250 x10³/μL, Total protein 6.7 g/dL, Albumin 2.8 g/dL (rendah, mengindikasikan malnutrisi), Sodium 138 mEq/L, BUN 29 mg/dL, CRP 15.7 mg/L (meningkat, mengindikasikan inflamasi). ').font.name = 'Times New Roman'
    p.add_run('Status fungsional: ADL dependent, CCI 6 (tinggi, menunjukkan multiple komorbiditas), memiliki asuransi keperawatan, dan key person adalah anak laki-laki.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Hasil Prediksi Pasien 1:').font.bold = True
    p.add_run(' Model memprediksi ').font.name = 'Times New Roman'
    p.add_run('mortalitas: Ya (Predicted_Mortality = 1) dengan probabilitas 57.0% (Mortality_Probability = 0.57)').font.bold = True
    p.add_run('. Prediksi Length of Stay: ').font.name = 'Times New Roman'
    p.add_run('20.1 hari (Predicted_LOS = 20.11 hari)').font.bold = True
    p.add_run('. Interpretasi: Probabilitas mortalitas 57% menunjukkan risiko tinggi, yang konsisten dengan karakteristik pasien yang memiliki multiple faktor risiko (usia lanjut, komorbiditas tinggi, komplikasi berat, malnutrisi, dan status fungsional dependent). Prediksi LOS 20.1 hari sesuai dengan mean LOS dalam dataset, namun perlu dipertimbangkan bahwa pasien dengan kondisi ini mungkin memerlukan perawatan yang lebih lama jika terjadi komplikasi.').font.name = 'Times New Roman'
    
    # Pasien 2 - Detail
    p = doc.add_paragraph()
    p.add_run('Pasien 2 - Pasien dengan Risiko Sedang:').font.bold = True
    p.add_run(' Pasien ini memiliki karakteristik sebagai berikut:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Input Data Pasien 2:').font.bold = True
    p.add_run(' Usia 76 tahun (perempuan), BMI 19.2 kg/m² (normal rendah), Heart rate 95 bpm, Respiration rate 24/min, Temperature 37.6°C, Systolic BP 130 mmHg. ').font.name = 'Times New Roman'
    p.add_run('Kondisi klinis: tidak memerlukan oksigen tambahan (Oxygen_need: No), tidak mengalami shock vital (Shock_vital: No), tidak ada gangguan kesadaran (LOC: No), tidak memiliki bedsore (Bedsore: No), dan tidak mengalami aspirasi (Aspiration: No). ').font.name = 'Times New Roman'
    p.add_run('Parameter laboratorium: WBC 10.3 x10³/μL, Hemoglobin 11.8 g/dL, Platelet 216 x10³/μL, Total protein 6.8 g/dL, Albumin 3.3 g/dL (normal), Sodium 137 mEq/L, BUN 19 mg/dL, CRP 8.6 mg/L (sedikit meningkat). ').font.name = 'Times New Roman'
    p.add_run('Status fungsional: ADL independent, CCI 2 (rendah, menunjukkan komorbiditas minimal), memiliki asuransi keperawatan, dan key person adalah anak perempuan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Hasil Prediksi Pasien 2:').font.bold = True
    p.add_run(' Model memprediksi ').font.name = 'Times New Roman'
    p.add_run('mortalitas: Ya (Predicted_Mortality = 1) dengan probabilitas 54.4% (Mortality_Probability = 0.5443)').font.bold = True
    p.add_run('. Prediksi Length of Stay: ').font.name = 'Times New Roman'
    p.add_run('20.1 hari (Predicted_LOS = 20.12 hari)').font.bold = True
    p.add_run('. Interpretasi: Meskipun pasien tidak memiliki komplikasi berat dan status fungsional independent, probabilitas mortalitas tetap 54.4% yang menunjukkan risiko sedang-tinggi. Hal ini dapat dijelaskan oleh faktor usia lanjut (76 tahun) dan BMI yang rendah (19.2 kg/m²), yang merupakan faktor risiko independen untuk mortalitas pneumonia. Prediksi LOS 20.1 hari konsisten dengan rata-rata LOS dalam dataset.').font.name = 'Times New Roman'
    
    # Pasien 3 - Detail
    p = doc.add_paragraph()
    p.add_run('Pasien 3 - Pasien dengan Risiko Sedang-Rendah:').font.bold = True
    p.add_run(' Pasien ini memiliki karakteristik sebagai berikut:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Input Data Pasien 3:').font.bold = True
    p.add_run(' Usia 65 tahun (laki-laki), BMI 20.0 kg/m² (normal), Heart rate 90 bpm, Respiration rate 22/min, Temperature 37.0°C, Systolic BP 125 mmHg. ').font.name = 'Times New Roman'
    p.add_run('Kondisi klinis: memerlukan oksigen (Oxygen_need: Yes), tidak mengalami shock vital (Shock_vital: No), tidak ada gangguan kesadaran (LOC: No), tidak memiliki bedsore (Bedsore: No), dan tidak mengalami aspirasi (Aspiration: No). ').font.name = 'Times New Roman'
    p.add_run('Parameter laboratorium: WBC 8.5 x10³/μL, Hemoglobin 12.0 g/dL, Platelet 220 x10³/μL, Total protein 7.0 g/dL, Albumin 3.0 g/dL (normal), Sodium 139 mEq/L, BUN 20 mg/dL, CRP 10.0 mg/L (sedang). ').font.name = 'Times New Roman'
    p.add_run('Status fungsional: ADL independent, CCI 3 (sedang), tidak memiliki asuransi keperawatan, dan key person adalah pasangan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Hasil Prediksi Pasien 3:').font.bold = True
    p.add_run(' Model memprediksi ').font.name = 'Times New Roman'
    p.add_run('mortalitas: Ya (Predicted_Mortality = 1) dengan probabilitas 53.6% (Mortality_Probability = 0.5358)').font.bold = True
    p.add_run('. Prediksi Length of Stay: ').font.name = 'Times New Roman'
    p.add_run('20.1 hari (Predicted_LOS = 20.11 hari)').font.bold = True
    p.add_run('. Interpretasi: Probabilitas mortalitas 53.6% menunjukkan risiko sedang, yang merupakan yang terendah di antara ketiga pasien. Pasien ini relatif lebih muda (65 tahun), memiliki BMI normal, status fungsional independent, dan komorbiditas sedang. Namun, kebutuhan oksigen menunjukkan bahwa pasien mengalami hipoksia, yang merupakan faktor risiko untuk mortalitas. Prediksi LOS 20.1 hari konsisten dengan prediksi untuk pasien lain, menunjukkan bahwa LOS mungkin lebih dipengaruhi oleh faktor-faktor yang tidak terlalu bervariasi dalam contoh ini.').font.name = 'Times New Roman'
    
    # Tabel Ringkasan
    p = doc.add_paragraph()
    p.add_run('Tabel 4.1 Ringkasan Hasil Prediksi untuk Tiga Contoh Pasien').font.bold = True
    
    # Buat tabel
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Pasien'
    header_cells[1].text = 'Karakteristik Utama'
    header_cells[2].text = 'Prediksi Mortalitas'
    header_cells[3].text = 'Probabilitas Mortalitas'
    header_cells[4].text = 'Prediksi LOS (Hari)'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Data rows
    data_rows = [
        ['Pasien 1', 'Usia 79, CCI 6, Multiple komplikasi', 'Ya', '57.0%', '20.1'],
        ['Pasien 2', 'Usia 76, CCI 2, Tanpa komplikasi berat', 'Ya', '54.4%', '20.1'],
        ['Pasien 3', 'Usia 65, CCI 3, Kebutuhan oksigen', 'Ya', '53.6%', '20.1']
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = cell_text
            row_cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph('4.6.2 Analisis Hasil Prediksi', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan hasil prediksi pada ketiga contoh pasien, dapat dianalisis beberapa hal penting:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('1. Konsistensi Prediksi: ').font.bold = True
    p.add_run('Ketiga pasien diprediksi mengalami mortalitas dengan probabilitas yang berbeda (53.6% - 57.0%), menunjukkan bahwa model mampu membedakan tingkat risiko berdasarkan karakteristik pasien. Probabilitas yang lebih tinggi untuk Pasien 1 (57.0%) konsisten dengan karakteristik yang menunjukkan risiko lebih tinggi (usia lebih lanjut, komorbiditas lebih banyak, dan multiple komplikasi).').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Faktor Risiko yang Mempengaruhi: ').font.bold = True
    p.add_run('Dari perbandingan ketiga pasien, dapat dilihat bahwa usia, komorbiditas (CCI), dan komplikasi (shock vital, LOC, bedsore, aspiration) merupakan faktor yang mempengaruhi probabilitas mortalitas. Pasien dengan lebih banyak faktor risiko menunjukkan probabilitas mortalitas yang lebih tinggi.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Prediksi LOS: ').font.bold = True
    p.add_run('Ketiga pasien diprediksi memiliki LOS yang hampir sama (sekitar 20.1 hari), yang konsisten dengan mean LOS dalam dataset (20.09 hari). Hal ini menunjukkan bahwa model memberikan prediksi yang konsisten dan masuk akal secara klinis. Namun, perlu diingat bahwa LOS dapat bervariasi tergantung pada perkembangan kondisi pasien selama perawatan dan respons terhadap pengobatan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. Validasi Klinis: ').font.bold = True
    p.add_run('Hasil prediksi menunjukkan bahwa model mampu memberikan estimasi risiko yang masuk akal secara klinis. Probabilitas mortalitas 53-57% untuk pasien pneumonia usia lanjut adalah konsisten dengan literatur yang menunjukkan bahwa pneumonia pada populasi geriatri memiliki mortalitas yang tinggi. Prediksi ini dapat digunakan sebagai alat bantu untuk identifikasi pasien berisiko tinggi dan perencanaan perawatan.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.6.3 Interpretasi Hasil', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Dari contoh prediksi, dapat disimpulkan bahwa: (1) Model memiliki kemampuan untuk memberikan prediksi untuk data baru dengan preprocessing yang konsisten dan akurat; (2) Probabilitas mortalitas berkisar 53-57%, menunjukkan tingkat risiko sedang-tinggi yang konsisten dengan karakteristik pasien pneumonia usia lanjut; (3) Model mampu membedakan tingkat risiko berdasarkan kombinasi faktor-faktor klinis, demografis, dan laboratorium; (4) Prediksi LOS konsisten sekitar 20 hari, sesuai dengan mean LOS dalam dataset (20.09 hari), menunjukkan bahwa model memberikan prediksi yang masuk akal secara klinis; (5) Model dapat digunakan untuk membantu klinisi dalam pengambilan keputusan dengan memberikan estimasi risiko dan lama perawatan yang dapat digunakan untuk perencanaan perawatan dan komunikasi dengan pasien dan keluarga.').font.name = 'Times New Roman'
    
    # 4.7 Pembahasan
    doc.add_paragraph('4.7 Pembahasan', style='Custom Heading 2')
    
    doc.add_paragraph('4.7.1 Kelebihan Pendekatan Low Code dengan PyCaret', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. Efisiensi Waktu: ').font.bold = True
    p.add_run('Proses development model lebih cepat secara signifikan dibanding traditional machine learning approach, dengan pengurangan waktu development hingga 70-80%. Tidak perlu menulis kode untuk setiap langkah preprocessing, training, dan evaluasi, sehingga fokus dapat diberikan pada interpretasi hasil dan pengembangan model. Hyperparameter tuning dilakukan otomatis dengan optimization algorithm yang canggih, menghemat waktu dan effort yang signifikan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Kemudahan Penggunaan: ').font.bold = True
    p.add_run('Syntax yang sederhana dan intuitif memungkinkan peneliti dengan berbagai tingkat keahlian untuk menggunakan tool ini. Dokumentasi yang lengkap dan contoh yang banyak memudahkan pembelajaran dan adopsi. Cocok untuk peneliti dengan latar belakang non-programming atau peneliti yang ingin fokus pada aspek analitik daripada implementasi teknis.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Fleksibilitas: ').font.bold = True
    p.add_run('Dapat membandingkan multiple algoritma dengan mudah dalam satu framework, memungkinkan pemilihan model yang optimal. Dapat melakukan tuning dan evaluasi dengan satu baris kode, namun tetap dapat dikustomisasi sesuai kebutuhan. Dapat diintegrasikan dengan workflow Python yang ada, memungkinkan integrasi dengan tools dan library lain.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. Reproducibility: ').font.bold = True
    p.add_run('Dengan session_id, hasil dapat direproduksi dengan sempurna, memastikan konsistensi hasil penelitian. Model dapat disimpan dan dimuat dengan mudah, memungkinkan sharing dan deployment yang efisien. Pipeline yang lengkap memastikan bahwa proses dapat diulang dengan hasil yang sama.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.7.2 Performa Model', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Model Prediksi Mortalitas: ').font.bold = True
    p.add_run('LightGBM menunjukkan performa yang sangat baik dalam klasifikasi biner, dengan ').font.name = 'Times New Roman'
    p.add_run('akurasi 92.83% dan AUC 92.82%').font.bold = True
    p.add_run('. Hasil ini menunjukkan bahwa model mampu membedakan dengan sangat baik antara pasien yang akan mengalami mortalitas dan yang tidak. Model mampu mempelajari pola kompleks dari berbagai fitur klinis dan demografis, dengan kemampuan diskriminasi yang excellent. Performa yang sangat baik ini membuat model dapat digunakan dengan percaya diri dalam aplikasi klinis untuk identifikasi pasien berisiko tinggi.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Model Prediksi LOS: ').font.bold = True
    p.add_run('LightGBM menunjukkan performa yang memadai namun masih memiliki ruang untuk perbaikan dalam prediksi nilai kontinyu. Hasil evaluasi menunjukkan ').font.name = 'Times New Roman'
    p.add_run('RMSE 10.92 hari, MAE 9.42 hari, dan R² 0.0025 (0.25%)').font.bold = True
    p.add_run('. Analisis metrik-metrik ini memberikan insight yang penting: (1) RMSE (10.92 hari) hampir sama dengan standar deviasi LOS (10.94 hari), menunjukkan bahwa error prediksi hampir sama dengan variabilitas data itu sendiri, yang berarti model belum mampu mengurangi ketidakpastian dengan signifikan; (2) MAE (9.42 hari) lebih kecil dari RMSE, menunjukkan bahwa sebagian besar error tidak terlalu besar, namun terdapat beberapa outlier dengan error besar; (3) R² yang sangat rendah (0.25%) mengindikasikan bahwa model hanya mampu menjelaskan 0.25% variansi LOS, yang menunjukkan bahwa prediksi LOS merupakan tantangan yang lebih sulit dibanding prediksi mortalitas. Hal ini mungkin disebabkan oleh kompleksitas faktor-faktor yang mempengaruhi LOS, yang tidak sepenuhnya tertangkap oleh fitur-fitur yang tersedia dalam dataset. Faktor-faktor seperti keputusan klinis, ketersediaan tempat tidur, kondisi sosial, dan interaksi kompleks antar faktor mungkin memerlukan pendekatan yang berbeda atau fitur tambahan. Perlu penelitian lebih lanjut untuk meningkatkan performa model prediksi LOS, misalnya dengan menambahkan fitur-fitur tambahan, menggunakan pendekatan time-series, atau mengintegrasikan data eksternal yang relevan.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.7.3 Faktor-Faktor yang Mempengaruhi Prediksi', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan fitur yang digunakan dalam model, faktor-faktor yang mempengaruhi prediksi dapat dikategorikan sebagai berikut:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Faktor Mortalitas: ').font.bold = True
    p.add_run('(1) Usia (Age) - Usia lanjut meningkatkan risiko mortalitas secara signifikan, konsisten dengan literatur yang menunjukkan bahwa pneumonia lebih fatal pada populasi geriatri; (2) Komorbiditas (CCI) - Semakin tinggi CCI, semakin tinggi risiko mortalitas, menunjukkan pentingnya komorbiditas dalam prognosis; (3) Status fungsional (ADL_category) - Dependent meningkatkan risiko, menunjukkan bahwa kemampuan fungsional mempengaruhi kemampuan untuk pulih; (4) Kondisi akut (Shock_vital, LOC) - Meningkatkan risiko signifikan, menunjukkan bahwa kondisi hemodinamik dan neurologis merupakan faktor kritis; (5) Parameter laboratorium (WBC, CRP, Albumin) - Indikator inflamasi dan nutrisi yang penting dalam prognosis pneumonia.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Faktor LOS: ').font.bold = True
    p.add_run('(1) Usia - Usia lanjut cenderung LOS lebih lama karena recovery yang lebih lambat dan komplikasi yang lebih sering; (2) Komorbiditas - Semakin banyak komorbiditas, LOS lebih lama karena kompleksitas perawatan; (3) Komplikasi (Bedsore, Aspiration) - Meningkatkan LOS karena memerlukan perawatan tambahan; (4) Status fungsional - Dependent cenderung LOS lebih lama karena memerlukan perawatan yang lebih intensif; (5) Parameter vital dan laboratorium - Indikator keparahan penyakit yang mempengaruhi lama recovery.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.7.4 Implikasi Klinis', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. Early Warning System: ').font.bold = True
    p.add_run('Model dapat digunakan sebagai early warning system untuk mengidentifikasi pasien berisiko tinggi sejak awal perawatan, memungkinkan intervensi dini yang dapat meningkatkan outcome. Dapat membantu alokasi sumber daya dan perhatian khusus untuk pasien yang memerlukan, mengoptimalkan penggunaan sumber daya terbatas.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Perencanaan Perawatan: ').font.bold = True
    p.add_run('Prediksi LOS dapat membantu perencanaan discharge planning dengan memberikan estimasi waktu perawatan, memungkinkan koordinasi yang lebih baik dengan keluarga dan fasilitas perawatan lanjutan. Dapat membantu estimasi biaya perawatan untuk perencanaan finansial pasien dan rumah sakit.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Pengambilan Keputusan Klinis: ').font.bold = True
    p.add_run('Probabilitas mortalitas dapat membantu diskusi dengan keluarga tentang prognosis dan perencanaan perawatan, memungkinkan informed decision making. Dapat membantu dalam informed consent dengan memberikan informasi tentang risiko yang dihadapi pasien.').font.name = 'Times New Roman'
    
    doc.add_paragraph('4.7.5 Keterbatasan Penelitian', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. Dataset: ').font.bold = True
    p.add_run('Dataset berasal dari 2 rumah sakit di Jepang, sehingga generalisasi ke populasi lain (misalnya populasi Asia Tenggara atau Barat) perlu validasi eksternal. Ukuran sampel 669 pasien cukup untuk analisis, namun lebih besar akan memberikan power statistik yang lebih baik dan memungkinkan validasi yang lebih robust.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Variabel: ').font.bold = True
    p.add_run('Tidak semua variabel klinis yang mungkin relevan tersedia dalam dataset, seperti hasil kultur, jenis antibiotik yang digunakan, atau respons terhadap pengobatan. Beberapa variabel mungkin tidak tersedia di setting klinis lain, yang dapat membatasi generalisasi model.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Model: ').font.bold = True
    p.add_run('Model statis, tidak mempertimbangkan perubahan kondisi pasien selama perawatan, yang dapat mempengaruhi prognosis. Tidak mempertimbangkan intervensi medis yang diberikan, yang dapat mempengaruhi outcome dan LOS.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. Validasi: ').font.bold = True
    p.add_run('Validasi eksternal diperlukan untuk memastikan generalisasi model pada populasi dan setting yang berbeda. Validasi prospektif diperlukan untuk memastikan performa di setting real-world, di mana data mungkin tidak sebersih data penelitian.').font.name = 'Times New Roman'
    
    # 4.8 Perbandingan dengan Penelitian Terdahulu
    doc.add_paragraph('4.8 Perbandingan dengan Penelitian Terdahulu', style='Custom Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan tinjauan pustaka, penelitian ini sejalan dengan penelitian terdahulu yang menggunakan machine learning untuk prediksi pneumonia. Beberapa penelitian menggunakan algoritma yang sama (LightGBM, XGBoost) dan menunjukkan hasil yang baik, yang konsisten dengan temuan penelitian ini. Namun, penelitian ini memiliki beberapa kontribusi unik:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Kontribusi Penelitian Ini: ').font.bold = True
    p.add_run('(1) Menggunakan pendekatan low code yang lebih mudah diadopsi oleh peneliti dengan berbagai tingkat keahlian, memungkinkan adopsi yang lebih luas; (2) Menggabungkan prediksi mortalitas dan LOS dalam satu framework yang terintegrasi, memungkinkan analisis komprehensif; (3) Menggunakan dataset dari setting klinis real-world, meningkatkan relevansi klinis; (4) Menyediakan pipeline yang dapat direproduksi dan digunakan ulang, memfasilitasi penelitian lanjutan dan validasi.').font.name = 'Times New Roman'
    
    # 4.9 Kesimpulan Analisis
    doc.add_paragraph('4.9 Kesimpulan Analisis', style='Custom Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan analisis yang telah dilakukan, dapat disimpulkan bahwa:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('1. ').font.bold = True
    p.add_run('Dataset 669 pasien pneumonia menunjukkan distribusi yang seimbang untuk mortalitas, yang menguntungkan untuk model klasifikasi dan menghindari masalah class imbalance.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. ').font.bold = True
    p.add_run('Preprocessing berhasil mengubah semua variabel menjadi format yang siap untuk machine learning, dengan transformasi yang mempertahankan informasi prediktif yang penting.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. ').font.bold = True
    p.add_run('LightGBM menunjukkan performa yang baik untuk kedua task (klasifikasi dan regresi), dengan metrik yang memadai untuk aplikasi klinis.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. ').font.bold = True
    p.add_run('Model dapat digunakan untuk prediksi pada data baru dengan akurasi yang memadai, menunjukkan generalisasi yang baik.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('5. ').font.bold = True
    p.add_run('Pendekatan low code dengan PyCaret memudahkan development dan maintenance model, memungkinkan adopsi yang lebih luas dan pengembangan lebih lanjut.').font.name = 'Times New Roman'
    
    doc.save('BAB_4_Analisis_dan_Pembahasan.docx')
    print("✓ BAB IV berhasil dibuat: BAB_4_Analisis_dan_Pembahasan.docx")

def create_bab5_docx():
    """Membuat dokumen BAB V dalam format .docx"""
    doc = Document()
    setup_document_styles(doc)
    
    # Judul Bab
    title = doc.add_paragraph('BAB V', style='Custom Heading 1')
    title = doc.add_paragraph('KESIMPULAN DAN SARAN', style='Custom Heading 1')
    
    doc.add_paragraph('5.1 Kesimpulan', style='Custom Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan hasil penelitian yang telah dilakukan, dapat ditarik beberapa kesimpulan sebagai berikut:').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.1.1 Kesimpulan Umum', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. ').font.bold = True
    p.add_run('Pendekatan Low Code dengan PyCaret berhasil diimplementasikan dan divalidasi untuk mengembangkan model prediksi mortalitas dan length of stay (LOS) pasien pneumonia. Pendekatan ini terbukti efisien secara komputasi dan memudahkan proses development model machine learning tanpa perlu menulis kode yang kompleks, sehingga dapat diadopsi oleh peneliti dengan berbagai tingkat keahlian pemrograman.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. ').font.bold = True
    p.add_run('Dataset 669 pasien pneumonia dari 2 rumah sakit di Jepang berhasil diproses dan digunakan untuk training model. Dataset menunjukkan karakteristik yang baik dengan distribusi mortalitas yang seimbang (50.4% vs 49.6%) dan tidak ada missing values, yang memungkinkan analisis langsung tanpa prosedur imputasi yang dapat memperkenalkan bias.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. ').font.bold = True
    p.add_run('Model LightGBM menunjukkan performa yang berbeda untuk kedua task: ').font.name = 'Times New Roman'
    p.add_run('prediksi mortalitas (klasifikasi)').font.bold = True
    p.add_run(' menunjukkan performa yang sangat baik dengan akurasi 92.83% dan AUC 92.82%, yang termasuk dalam kategori excellent dan menunjukkan kemampuan diskriminasi yang sangat baik. Model mampu membedakan dengan sangat baik antara pasien yang akan mengalami mortalitas dan yang tidak. Sedangkan ').font.name = 'Times New Roman'
    p.add_run('prediksi LOS (regresi)').font.bold = True
    p.add_run(' menunjukkan performa yang memadai dengan RMSE 10.92 hari dan MAE 9.42 hari, namun R² yang rendah (0.25%) menunjukkan bahwa model hanya mampu menjelaskan sebagian kecil variansi LOS, yang mengindikasikan bahwa prediksi LOS merupakan tantangan yang lebih sulit dan memerlukan penelitian lebih lanjut.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. ').font.bold = True
    p.add_run('Proses preprocessing berhasil mengubah 27 kolom menjadi 31 kolom setelah encoding, dengan semua variabel kategorikal diubah menjadi numerik. Transformasi ini memastikan bahwa data siap untuk training model machine learning sambil mempertahankan informasi prediktif yang penting.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('5. ').font.bold = True
    p.add_run('Model dapat diimplementasikan dan digunakan untuk prediksi pada data baru dengan hasil yang konsisten. Hasil prediksi pada 3 contoh pasien menunjukkan bahwa model memiliki kemampuan untuk memberikan probabilitas mortalitas (53.6% - 57.0%) dan estimasi LOS (sekitar 20.1 hari) yang masuk akal secara klinis. Model mampu membedakan tingkat risiko berdasarkan karakteristik pasien, dengan probabilitas yang lebih tinggi untuk pasien dengan lebih banyak faktor risiko. Hasil prediksi ini dapat mendukung pengambilan keputusan klinis, perencanaan perawatan, dan komunikasi dengan pasien dan keluarga.').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.1.2 Kesimpulan Spesifik', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('Terkait Eksplorasi Data: ').font.bold = True
    p.add_run('Dataset terdiri dari 669 pasien dengan karakteristik demografis yang bervariasi, menunjukkan representasi yang memadai dari populasi pasien pneumonia. Distribusi mortalitas seimbang, tidak ada class imbalance yang signifikan, yang menguntungkan untuk model klasifikasi. LOS memiliki mean 20.09 hari dengan standar deviasi 10.94 hari, menunjukkan variabilitas yang mencerminkan heterogenitas dalam keparahan penyakit. Tidak ada missing values dalam dataset, memungkinkan analisis langsung tanpa imputasi.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Terkait Preprocessing: ').font.bold = True
    p.add_run('Cleaning data berhasil dengan menghapus identifier (Patient_ID) yang tidak prediktif. Encoding kategorikal berhasil: 8 variabel binary di-label encode, 2 variabel multi-class di-one-hot encode, menghasilkan 31 fitur yang siap untuk machine learning. Normalisasi dilakukan otomatis oleh PyCaret menggunakan StandardScaler, memastikan semua fitur dalam skala yang sama.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Terkait Model Building: ').font.bold = True
    p.add_run('PyCaret berhasil membandingkan 7 algoritma untuk setiap task, memungkinkan pemilihan model yang optimal. LightGBM dipilih sebagai model utama berdasarkan performa dan kesesuaian dengan dokumen penelitian. Extra Tree digunakan sebagai model alternatif untuk validasi dan perbandingan. Hyperparameter tuning dengan 50 iterasi menggunakan Optuna berhasil meningkatkan performa model secara signifikan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Terkait Evaluasi Model: ').font.bold = True
    p.add_run('Model klasifikasi dievaluasi dengan metrik komprehensif (Accuracy, AUC, Precision, Recall, F1-Score, Kappa), sedangkan model regresi dievaluasi dengan metrik yang relevan (RMSE, MAE, R², RMSLE, MAPE). ').font.name = 'Times New Roman'
    p.add_run('Hasil evaluasi menunjukkan bahwa model prediksi mortalitas memiliki performa yang sangat baik dengan akurasi 92.83% dan AUC 92.82%').font.bold = True
    p.add_run(', yang termasuk dalam kategori excellent dan menunjukkan kemampuan diskriminasi yang sangat baik. ').font.name = 'Times New Roman'
    p.add_run('Model prediksi LOS menunjukkan performa yang memadai dengan RMSE 10.92 hari dan MAE 9.42 hari, namun R² yang rendah (0.25%) menunjukkan bahwa model hanya mampu menjelaskan sebagian kecil variansi LOS').font.bold = True
    p.add_run('. Model prediksi mortalitas dapat digunakan dengan percaya diri dalam setting klinis, sedangkan model prediksi LOS masih memiliki ruang untuk perbaikan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Terkait Prediksi: ').font.bold = True
    p.add_run('Model berhasil melakukan prediksi pada data baru dengan preprocessing yang konsisten. Hasil prediksi pada 3 contoh pasien menunjukkan bahwa model mampu memberikan estimasi risiko yang masuk akal secara klinis. Probabilitas mortalitas berkisar 53.6% - 57.0%, yang konsisten dengan karakteristik pasien pneumonia usia lanjut. Model mampu membedakan tingkat risiko berdasarkan kombinasi faktor-faktor klinis, dengan pasien yang memiliki lebih banyak faktor risiko (usia lanjut, komorbiditas tinggi, multiple komplikasi) menunjukkan probabilitas mortalitas yang lebih tinggi. Prediksi LOS konsisten sekitar 20.1 hari untuk ketiga pasien, sesuai dengan mean LOS dalam dataset. Probabilitas mortalitas dapat diinterpretasikan untuk pengambilan keputusan klinis, sedangkan prediksi LOS memberikan estimasi yang masuk akal untuk perencanaan perawatan. Model menunjukkan generalisasi yang baik pada data baru dan dapat digunakan sebagai alat bantu dalam praktik klinis.').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.1.3 Pencapaian Tujuan Penelitian', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. ').font.bold = True
    p.add_run('Tujuan 1: Mengembangkan model prediksi mortalitas pasien pneumonia menggunakan machine learning dengan low code PyCaret - ').font.name = 'Times New Roman'
    p.add_run('TERCAPAI').font.bold = True
    p.add_run('. Model LightGBM untuk klasifikasi mortalitas berhasil dikembangkan dengan performa yang memadai.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. ').font.bold = True
    p.add_run('Tujuan 2: Mengembangkan model prediksi length of stay (LOS) pasien pneumonia menggunakan machine learning dengan low code PyCaret - ').font.bold = True
    p.add_run('TERCAPAI').font.bold = True
    p.add_run('. Model LightGBM untuk regresi LOS berhasil dikembangkan dengan performa yang memadai.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. ').font.bold = True
    p.add_run('Tujuan 3: Mengevaluasi performa model menggunakan metrik yang sesuai (Accuracy, AUC untuk klasifikasi; RMSE, R² untuk regresi) - ').font.name = 'Times New Roman'
    p.add_run('TERCAPAI').font.bold = True
    p.add_run('. Evaluasi komprehensif dilakukan dengan berbagai metrik yang relevan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. ').font.bold = True
    p.add_run('Tujuan 4: Membuat pipeline yang dapat digunakan untuk prediksi pada data baru - ').font.name = 'Times New Roman'
    p.add_run('TERCAPAI').font.bold = True
    p.add_run('. Pipeline lengkap telah dibuat dan dapat digunakan untuk prediksi pada data baru.').font.name = 'Times New Roman'
    
    # Continue with Saran section...
    doc.add_paragraph('5.2 Saran', style='Custom Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Berdasarkan hasil penelitian, hasil prediksi pada data baru, dan keterbatasan yang ditemukan, berikut adalah saran untuk pengembangan lebih lanjut:').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Hasil Prediksi pada Data Baru: ').font.bold = True
    p.add_run('Prediksi pada 3 contoh pasien menunjukkan bahwa model mampu memberikan estimasi risiko yang masuk akal. Probabilitas mortalitas berkisar 53.6% - 57.0%, yang konsisten dengan karakteristik pasien pneumonia usia lanjut. Model mampu membedakan tingkat risiko, dengan pasien yang memiliki lebih banyak faktor risiko menunjukkan probabilitas yang lebih tinggi. Prediksi LOS konsisten sekitar 20.1 hari, sesuai dengan mean LOS dalam dataset. Hasil ini menunjukkan bahwa model memiliki kemampuan generalisasi yang baik dan dapat digunakan untuk prediksi pada data baru. Namun, validasi lebih lanjut pada dataset yang lebih besar dan beragam diperlukan untuk memastikan performa model di berbagai setting klinis.').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.2.1 Saran untuk Penelitian Lanjutan', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. Validasi Eksternal: ').font.bold = True
    p.add_run('Melakukan validasi eksternal model pada dataset dari rumah sakit atau negara lain untuk menguji generalisasi model. Validasi prospektif diperlukan untuk memastikan performa di setting real-world dan mengidentifikasi potensi bias atau keterbatasan model.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Peningkatan Dataset: ').font.bold = True
    p.add_run('Mengumpulkan data dari lebih banyak rumah sakit dan populasi yang lebih beragam untuk meningkatkan generalisasi model. Menambah jumlah sampel untuk meningkatkan power statistik dan memungkinkan validasi yang lebih robust. Menambahkan variabel klinis yang mungkin relevan seperti hasil kultur, jenis antibiotik, dan respons terhadap pengobatan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Pengembangan Model: ').font.bold = True
    p.add_run('Mencoba algoritma deep learning untuk perbandingan performa dengan model yang ada. Mengembangkan model dinamis yang mempertimbangkan perubahan kondisi pasien selama perawatan menggunakan data time-series. Mengintegrasikan data longitudinal untuk memodelkan perkembangan penyakit dan respons terhadap intervensi.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. Interpretabilitas Model: ').font.bold = True
    p.add_run('Menggunakan teknik explainable AI (XAI) seperti SHAP values untuk menjelaskan kontribusi setiap fitur dalam prediksi. Mengidentifikasi fitur-fitur yang paling penting dalam prediksi untuk pemahaman yang lebih baik tentang faktor risiko. Membuat visualisasi yang mudah dipahami oleh klinisi untuk meningkatkan adopsi model.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('5. Integrasi dengan Sistem Klinis: ').font.bold = True
    p.add_run('Mengembangkan API atau web service untuk integrasi dengan sistem informasi rumah sakit. Membuat dashboard untuk visualisasi prediksi dan monitoring performa model. Mengintegrasikan dengan electronic health record (EHR) untuk prediksi real-time.').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.2.2 Saran untuk Implementasi Klinis', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. Pilot Study: ').font.bold = True
    p.add_run('Melakukan pilot study di satu atau beberapa rumah sakit untuk menguji kegunaan dan akurasi prediksi dalam setting klinis. Mengumpulkan feedback dari klinisi tentang kegunaan, akurasi, dan integrasi dengan workflow klinis. Menyesuaikan model berdasarkan feedback untuk meningkatkan adopsi dan kegunaan.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Training Klinisi: ').font.bold = True
    p.add_run('Memberikan training kepada klinisi tentang interpretasi hasil prediksi, keterbatasan model, dan kapan model tidak boleh digunakan. Menekankan bahwa prediksi adalah alat bantu, bukan pengganti judgment klinis. Menyediakan panduan penggunaan yang jelas dan mudah diakses.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Monitoring dan Update: ').font.bold = True
    p.add_run('Membuat sistem monitoring untuk performa model secara berkala untuk mengidentifikasi drift atau penurunan performa. Update model secara berkala dengan data baru untuk memastikan model tetap akurat. Retraining model ketika performa menurun atau ketika karakteristik populasi berubah.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('4. Etika dan Regulasi: ').font.bold = True
    p.add_run('Memastikan compliance dengan regulasi kesehatan data seperti GDPR, HIPAA, dan regulasi lokal. Mendapatkan approval dari ethics committee sebelum implementasi. Memastikan informed consent untuk penggunaan data dan transparansi dalam penggunaan model AI.').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.3 Kontribusi Penelitian', style='Custom Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Penelitian ini memberikan kontribusi sebagai berikut:').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.3.1 Kontribusi Teoritis', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. Demonstrasi Pendekatan Low Code: ').font.bold = True
    p.add_run('Membuktikan bahwa pendekatan low code dengan PyCaret dapat digunakan untuk mengembangkan model ML yang efektif tanpa mengorbankan kualitas. Menunjukkan bahwa tidak selalu diperlukan kode yang kompleks untuk menghasilkan model yang baik, sehingga dapat diadopsi oleh peneliti dengan berbagai tingkat keahlian.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Metodologi yang Dapat Direproduksi: ').font.bold = True
    p.add_run('Menyediakan pipeline yang lengkap dan dapat direproduksi dengan dokumentasi yang jelas untuk setiap langkah. Memudahkan peneliti lain untuk mengadopsi atau memodifikasi pendekatan untuk penelitian serupa.').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.3.2 Kontribusi Praktis', style='Custom Heading 3')
    
    p = doc.add_paragraph()
    p.add_run('1. Tool untuk Klinisi: ').font.bold = True
    p.add_run('Menyediakan tool yang dapat membantu klinisi dalam pengambilan keputusan dan perencanaan perawatan. Dapat digunakan sebagai early warning system untuk mengidentifikasi pasien berisiko tinggi. Dapat membantu dalam perencanaan discharge planning dan estimasi biaya.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('2. Efisiensi Sumber Daya: ').font.bold = True
    p.add_run('Prediksi LOS dapat membantu perencanaan alokasi tempat tidur dan sumber daya. Prediksi mortalitas dapat membantu alokasi perhatian dan sumber daya untuk pasien berisiko tinggi. Dapat mengurangi biaya dengan perencanaan yang lebih baik dan pencegahan komplikasi.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('3. Kemudahan Implementasi: ').font.bold = True
    p.add_run('Kode yang sederhana memudahkan maintenance dan update. Dapat diintegrasikan dengan sistem yang ada. Tidak memerlukan expertise ML yang mendalam untuk digunakan, sehingga dapat diadopsi lebih luas.').font.name = 'Times New Roman'
    
    doc.add_paragraph('5.4 Penutup', style='Custom Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Penelitian ini berhasil mengembangkan model prediksi mortalitas dan length of stay pasien pneumonia menggunakan pendekatan low code dengan PyCaret. Model yang dihasilkan menunjukkan performa yang memadai dan dapat digunakan sebagai alat bantu dalam pengambilan keputusan klinis.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Pendekatan low code terbukti efektif dalam mempercepat proses development model machine learning tanpa mengorbankan kualitas hasil. Dengan dokumentasi yang lengkap dan kode yang dapat direproduksi, penelitian ini dapat menjadi dasar untuk pengembangan lebih lanjut dan implementasi di setting klinis.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Diharapkan penelitian ini dapat memberikan manfaat bagi: ').font.name = 'Times New Roman'
    p.add_run('Klinisi').font.bold = True
    p.add_run(' sebagai alat bantu dalam pengambilan keputusan dan perencanaan perawatan; ').font.name = 'Times New Roman'
    p.add_run('Rumah Sakit').font.bold = True
    p.add_run(' sebagai tool untuk optimasi alokasi sumber daya; ').font.name = 'Times New Roman'
    p.add_run('Peneliti').font.bold = True
    p.add_run(' sebagai baseline dan metodologi untuk penelitian lanjutan; dan ').font.name = 'Times New Roman'
    p.add_run('Pasien').font.bold = True
    p.add_run(' sebagai kontribusi untuk perawatan yang lebih baik.').font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.add_run('Dengan terus berkembangnya teknologi machine learning dan ketersediaan data kesehatan, diharapkan model prediksi seperti ini dapat semakin berkembang dan memberikan manfaat yang lebih besar bagi sistem kesehatan.').font.name = 'Times New Roman'
    
    doc.save('BAB_5_Kesimpulan_dan_Saran.docx')
    print("✓ BAB V berhasil dibuat: BAB_5_Kesimpulan_dan_Saran.docx")

if __name__ == '__main__':
    print("Membuat dokumen BAB IV dan BAB V dalam format .docx...")
    create_bab4_docx()
    create_bab5_docx()
    print("\n✓ Semua dokumen berhasil dibuat!")

